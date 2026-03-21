from http.server import BaseHTTPRequestHandler
import json
import io
import base64
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

class handler(BaseHTTPRequestHandler):
    def do_POST(self):
        content_length = int(self.headers['Content-Length'])
        body = self.rfile.read(content_length)
        data = json.loads(body)

        minutes = data.get('annotated_minutes', '')
        meeting_name = data.get('meeting_name', 'Meeting')
        project = data.get('project', '')
        date = data.get('date', '')
        attendees = data.get('attendees', '')

        doc = Document()

        # Page margins — 1 inch all sides
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Default font
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)

        # HEADER SECTION
        def add_para(text, bold=False, size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
            p = doc.add_paragraph()
            p.alignment = align
            run = p.add_run(text)
            run.bold = bold
            run.font.name = 'Arial'
            run.font.size = Pt(size)
            return p

        add_para('MEETING MINUTES', bold=True, size=11)
        add_para(project, bold=True, size=16)
        add_para(meeting_name, bold=True, size=16)
        add_para(date, bold=True, size=16)
        doc.add_paragraph()

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run1 = p.add_run('Attendees: ')
        run1.bold = True
        run1.font.name = 'Arial'
        run1.font.size = Pt(11)
        run2 = p.add_run(attendees)
        run2.font.name = 'Arial'
        run2.font.size = Pt(11)
        doc.add_paragraph()

        # TABLE
        # Parse annotated minutes into rows
        lines = minutes.split('\n')
        rows = []
        current_num = ''
        current_item = []
        current_notes = []
        current_action = ''

        def is_num_line(l):
            return bool(re.match(r'^\d+\.\d+', l.strip()))

        def is_section_line(l):
            return bool(re.match(r'^\d+\.0', l.strip()))

        def is_notes_line(l):
            return 'notes:' in l.lower() or '<strong>notes' in l.lower()

        def is_action_line(l):
            return '\u2192' in l or ('->' in l and re.match(r'^[A-Z]', l.strip()))

        def is_person_name(l):
            # Standalone person name: 2–4 title-case words, short line, no digits/punctuation
            # e.g. "Aaron Cutler", "Juan Anderson", "Mary Jo Smith"
            t = l.strip()
            return bool(re.match(r'^[A-Z][A-Za-z\-]+(?: [A-Z][A-Za-z\-]+){1,3}$', t)) and len(t) < 50

        def flush_row():
            if not current_num:
                return
            rows.append({
                'num': current_num,
                'item': list(current_item),
                'notes': list(current_notes),
                'action': current_action,
                'is_section': is_section_line(current_num + ' ')
            })

        for line in lines:
            t = line.strip()
            if not t:
                continue
            if is_num_line(t):
                flush_row()
                current_num = ''
                current_item = []
                current_notes = []
                current_action = ''
                match = re.match(r'^(\d+\.\d+)\s*(.*)', t)
                if match:
                    current_num = match.group(1)
                    if match.group(2):
                        current_item.append(match.group(2))
            elif is_notes_line(t) or is_action_line(t):
                current_notes.append(t)
            elif current_num and not current_notes and is_person_name(t):
                # Standalone person name appearing before Notes: goes in the Action column
                current_action = t
            elif current_num:
                if not current_notes:
                    current_item.append(t)
                else:
                    current_notes.append(t)

        flush_row()

        def clean_html(text):
            text = re.sub(r'<strong>(.*?)</strong>', r'\1', text)
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            return text.strip()

        def add_bold_run(para, text):
            pattern = r'(<strong>.*?</strong>|\*\*.*?\*\*)'
            parts = re.split(pattern, text)
            for part in parts:
                if re.match(r'<strong>(.*?)</strong>', part):
                    content = re.sub(r'<strong>(.*?)</strong>', r'\1', part)
                    run = para.add_run(content)
                    run.bold = True
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                elif re.match(r'\*\*(.+?)\*\*', part):
                    content = re.sub(r'\*\*(.+?)\*\*', r'\1', part)
                    run = para.add_run(content)
                    run.bold = True
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                elif part:
                    run = para.add_run(part)
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)

        # Create table with header
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'

        # Col widths: 0.6in (No.) + 4.7in (Item) + 1.2in (Action) = 6.5in
        widths = [Inches(0.6), Inches(4.7), Inches(1.2)]
        for i, width in enumerate(widths):
            for cell in table.columns[i].cells:
                cell.width = width

        # Header row
        hdr = table.rows[0]
        hdr_labels = ['No.', 'Item', 'Action']
        for i, label in enumerate(hdr_labels):
            cell = hdr.cells[i]
            cell.paragraphs[0].clear()
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(label)
            run.bold = True
            run.font.name = 'Arial'
            run.font.size = Pt(11)

        # Data rows
        for row_data in rows:
            row = table.add_row()

            # Col 1 — Item number
            cell0 = row.cells[0]
            cell0.paragraphs[0].clear()
            p0 = cell0.paragraphs[0]
            run = p0.add_run(row_data['num'])
            run.bold = row_data['is_section']
            run.font.name = 'Arial'
            run.font.size = Pt(11)

            # Col 2 — Item text + notes
            cell1 = row.cells[1]
            cell1.paragraphs[0].clear()

            for i, item_line in enumerate(row_data['item']):
                if i == 0:
                    p = cell1.paragraphs[0]
                else:
                    p = cell1.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                run = p.add_run(item_line)
                run.bold = row_data['is_section']
                run.font.name = 'Arial'
                run.font.size = Pt(11)

            for note_line in row_data['notes']:
                p = cell1.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.left_indent = Inches(0.3)
                if 'notes:' in note_line.lower() or '<strong>notes' in note_line.lower():
                    note_text = re.sub(
                        r'(<strong>)?[Nn]otes:(</strong>)?',
                        '', note_line
                    ).strip()
                    label_run = p.add_run('Notes: ')
                    label_run.bold = True
                    label_run.font.name = 'Arial'
                    label_run.font.size = Pt(11)
                    add_bold_run(p, note_text)
                elif '\u2192' in note_line or '->' in note_line:
                    sep = '\u2192' if '\u2192' in note_line else '->'
                    parts = note_line.split(sep, 1)
                    person = clean_html(parts[0]).strip()
                    action = clean_html(parts[1]).strip() if len(parts) > 1 else ''
                    person_run = p.add_run(person + ' \u2192 ')
                    person_run.bold = True
                    person_run.font.name = 'Arial'
                    person_run.font.size = Pt(11)
                    action_run = p.add_run(action)
                    action_run.font.name = 'Arial'
                    action_run.font.size = Pt(11)
                else:
                    add_bold_run(p, note_line)

            # Col 3 — Action person
            cell2 = row.cells[2]
            cell2.paragraphs[0].clear()
            p2 = cell2.paragraphs[0]
            p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p2.add_run(row_data['action'])
            run.font.name = 'Arial'
            run.font.size = Pt(11)

        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        doc_bytes = buffer.read()
        doc_base64 = base64.b64encode(doc_bytes).decode('utf-8')

        filename = f"{meeting_name}_MeetingMinutes_{date}.docx"
        response = json.dumps({
            'docx_base64': doc_base64,
            'filename': filename
        })

        self.send_response(200)
        self.send_header('Content-Type', 'application/json')
        self.end_headers()
        self.wfile.write(response.encode('utf-8'))
